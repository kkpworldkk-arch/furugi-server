from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ===== ページ余白 =====
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ===== スタイル定数 =====
BROWN      = RGBColor(0x5D, 0x40, 0x37)
BROWN_DARK = RGBColor(0x3E, 0x27, 0x23)
GRAY_TEXT  = RGBColor(0x55, 0x55, 0x55)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(text, level=1, color=BROWN_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(16 if level == 1 else 13)
    if level == 1:
        p.paragraph_format.border_bottom = True
    return p

def body(text, color=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def add_table(headers, rows, header_color='5D4037'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # データ行
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if r_idx % 2 == 0:
                set_cell_bg(cell, 'FAF6F3')
    return table

# ============================================================
#  表紙
# ============================================================
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('古着屋マップ')
title_run.bold = True
title_run.font.size = Pt(32)
title_run.font.color.rgb = BROWN_DARK

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run('事業企画書')
sub_run.font.size = Pt(18)
sub_run.font.color.rgb = BROWN

doc.add_paragraph()

tagline_p = doc.add_paragraph()
tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_run = tagline_p.add_run('「全国の古着屋を、地図でつなぐ。」')
tag_run.italic = True
tag_run.font.size = Pt(13)
tag_run.font.color.rgb = GRAY_TEXT

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f'{datetime.date.today().year}年{datetime.date.today().month}月　作成\nConfidential')
date_run.font.size = Pt(11)
date_run.font.color.rgb = GRAY_TEXT

doc.add_page_break()

# ============================================================
#  1. エグゼクティブサマリー
# ============================================================
heading('1. エグゼクティブサマリー')
body('「古着屋マップ」は、日本全国の古着屋を地図上で発見・評価・共有できる、古着特化型プラットフォームです。')
body('古着市場は国内で年率10%超の成長を続けているにもかかわらず、古着屋に特化した情報インフラは存在しません。'
     'Googleマップは汎用的すぎてニーズに応えられず、SNSは情報が散在して検索性が低い状態です。')
body('本サービスはその空白を埋め、古着ユーザー × 古着屋のマッチングプラットフォームとして、'
     '広告・店舗課金・データ提供の3軸で収益化を目指します。')
doc.add_paragraph()

add_table(
    ['項目', '内容'],
    [
        ['サービス名', '古着屋マップ'],
        ['形態', 'Webアプリ・モバイルアプリ（iOS / Android）'],
        ['現在の登録店舗数', '75店舗（2026年6月時点）'],
        ['ターゲットユーザー', '古着好きの10〜30代'],
        ['対象エリア', '日本全国（現在：関東・東海・東北中心）'],
    ]
)

doc.add_page_break()

# ============================================================
#  2. 課題と解決策
# ============================================================
heading('2. 課題と解決策')

heading('2-1. ユーザー側の課題', level=2)
body('・自分の好きなジャンルの古着屋がどこにあるかわからない')
body('・Googleマップでは古着屋をジャンルで絞り込む手段がない')
body('・SNSの情報は散らばっており、まとめて探せない')

heading('2-2. 店舗側の課題', level=2)
body('・小規模店が多く、広告費をかけられない')
body('・自店のジャンルや特徴を適切なターゲットに届ける手段が少ない')
body('・SNS運用で手一杯で、新規集客に課題を抱えている')

heading('2-3. 解決策', level=2)
body('古着屋マップ = 古着に特化した「ジャンル別・地図検索」プラットフォーム')
doc.add_paragraph()

add_table(
    ['機能', '効果'],
    [
        ['ジャンル別マップ検索', 'ヴィンテージ・アメカジ・ストリート等を即絞り込み'],
        ['口コミ・評価', '古着好きの視点に特化したレビュー'],
        ['お知らせ機能', 'セール・入荷情報を店舗からユーザーへ直接配信'],
        ['「行った！」記録', '訪問記録でリピート・新規開拓を促進'],
        ['ユーザー参加型登録', '利用者が店舗情報を追加・更新し情報精度を維持'],
    ]
)

doc.add_page_break()

# ============================================================
#  3. 市場規模
# ============================================================
heading('3. 市場規模')

heading('3-1. 国内古着・リユース市場', level=2)
add_table(
    ['区分', '規模', '成長率'],
    [
        ['国内リユース市場全体（2023年）', '約2兆9,000億円', '前年比 +8.4%'],
        ['うち古着・アパレル分野', '約9,000億円', '前年比 +12%'],
        ['国内古着屋店舗数（推計）', '約15,000〜20,000店', '増加傾向'],
    ]
)
doc.add_paragraph()
body('出典参考：環境省リユース市場規模調査、矢野経済研究所', color=GRAY_TEXT)

heading('3-2. TAM / SAM / SOM', level=2)
add_table(
    ['区分', '定義', '規模'],
    [
        ['TAM（全体市場）', '古着・ファッションに関心を持つ日本の10〜40代', '約2,800万人'],
        ['SAM（獲得可能市場）', 'スマホで古着屋を能動的に検索するユーザー', '約500万人'],
        ['SOM（3年以内の目標）', '月間アクティブユーザー（MAU）', '50万人'],
    ]
)

doc.add_page_break()

# ============================================================
#  4. プロダクト概要
# ============================================================
heading('4. プロダクト概要')

heading('4-1. 主要機能', level=2)
add_table(
    ['機能', '概要'],
    [
        ['地図検索', 'Google Maps連携。現在地周辺の古着屋をワンタップで表示'],
        ['ジャンル絞り込み', 'ヴィンテージ・アメカジ・ストリート・レディース・ブランド古着 等'],
        ['店舗詳細', '営業時間・定休日・駐車場・支払い方法・最寄り駅'],
        ['口コミ・評価', 'ユーザーが星評価とコメントを投稿'],
        ['お知らせ', '店舗からユーザーへのセール・入荷情報を配信'],
        ['行った！ボタン', '訪問記録を端末に保存。巡った店舗を一覧管理'],
        ['お気に入り登録', '再訪したい店舗をリスト化'],
        ['右側パネル', '現在の地図範囲に映っている店舗を即リスト表示（Googleマップ風）'],
    ]
)

heading('4-2. 技術スタック', level=2)
add_table(
    ['領域', '技術・サービス'],
    [
        ['フロントエンド', 'Flutter（iOS・Android・Web 1コードで3対応）'],
        ['バックエンド', 'Python Flask + PostgreSQL'],
        ['インフラ', 'Railway（API）、Netlify（Web）'],
        ['地図', 'Google Maps Flutter Plugin'],
    ]
)

doc.add_page_break()

# ============================================================
#  5. ビジネスモデル
# ============================================================
heading('5. ビジネスモデル')

heading('フェーズ1（〜1万DAU）：ユーザー獲得優先', level=2)
body('全機能を無料開放。SNS・ストリート配布等でユーザー数を積み上げる。')

heading('フェーズ2（1万〜10万DAU）：店舗向け課金スタート', level=2)
add_table(
    ['収益源', '概要', '想定単価'],
    [
        ['プレミアム店舗掲載', '上位表示・バナー・お知らせ配信', '月額3,000〜10,000円/店'],
        ['セール・イベント告知', '期間限定でユーザーに告知配信', '1回2,000〜5,000円'],
        ['データ提供（B2B）', 'エリア別・ジャンル別の集客データ', '要相談'],
    ]
)

heading('フェーズ3（10万DAU〜）：プラットフォーム化', level=2)
add_table(
    ['収益源', '概要'],
    [
        ['ユーザープレミアム', '高度な絞り込み・訪問履歴エクスポート等　月額300円'],
        ['ECアフィリエイト', 'メルカリ・ZOZOUSEDとの送客連携'],
        ['古着屋開業支援', '新規出店コンサルティング・掲載支援パッケージ'],
    ]
)

heading('3年後収益シミュレーション', level=2)
add_table(
    ['指標', '目標値'],
    [
        ['MAU', '50万人'],
        ['登録店舗数', '5,000店'],
        ['プレミアム契約店舗', '500店（契約率10%）'],
        ['店舗課金ARR', '約3,000万円/年'],
        ['広告・データ収益', '約1,500万円/年'],
        ['合計ARR', '約4,500万円/年'],
    ]
)

doc.add_page_break()

# ============================================================
#  6. 競合分析
# ============================================================
heading('6. 競合分析')
add_table(
    ['サービス', '強み', '弱み'],
    [
        ['Googleマップ', '認知度・網羅性', '古着特化でない・ジャンル絞り込み不可'],
        ['Instagram', 'ビジュアル・拡散力', '店舗を横断して検索できない'],
        ['古着note・ブログ', '情報の深さ', 'リアルタイム性なし・地図連携なし'],
        ['セカスト公式サイト', 'チェーン店情報', '独立系の古着屋を網羅していない'],
        ['古着屋マップ（当社）', '古着特化・地図・口コミ・告知の一体化', '現状、認知度が低い'],
    ]
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('▶ 直接競合となるサービスは現在存在しない。今が参入のゴールデンタイム。')
run.bold = True
run.font.color.rgb = BROWN

doc.add_page_break()

# ============================================================
#  7. トラクション
# ============================================================
heading('7. トラクション（現状の実績）')
add_table(
    ['項目', '内容'],
    [
        ['登録店舗数', '75店舗（関東・東海・東北中心）'],
        ['対応プラットフォーム', 'Web・iOS・Android（Flutter製）'],
        ['バックエンド', 'Railway上で24時間稼働中'],
        ['開発体制', '学生1名で設計・実装・運用を完結'],
        ['リリース済み機能', '口コミ・お知らせ・訪問記録・お気に入り・ピン位置調整 等'],
    ]
)

doc.add_page_break()

# ============================================================
#  8. ロードマップ
# ============================================================
heading('8. ロードマップ')
add_table(
    ['時期', 'マイルストーン'],
    [
        ['2026年 Q3（現在）', 'SNS告知開始 / 下北沢・高円寺エリア重点登録 / アプリストア申請'],
        ['2026年 Q4', '登録店舗300店達成 / 店舗ダッシュボード開発 / 写真投稿機能追加'],
        ['2027年 Q1〜Q2', 'プレミアム掲載プラン販売開始 / 全国主要都市展開 / MAU 1万人'],
        ['2027年 Q3〜', 'EC連携・アフィリエイト / MAU 10万人 / ARR 1,000万円達成'],
    ]
)

doc.add_page_break()

# ============================================================
#  9. 資金調達計画
# ============================================================
heading('9. 資金調達計画')

heading('調達希望額：●●万円', level=2)

heading('使途内訳', level=2)
add_table(
    ['用途', '比率', '具体例'],
    [
        ['マーケティング・PR', '50%', 'SNS広告・インフルエンサー施策・ストリートイベント'],
        ['開発強化', '30%', '写真投稿・店舗ダッシュボード・アプリストア申請'],
        ['店舗開拓営業', '15%', '下北沢・高円寺等の集中営業・掲載交渉'],
        ['運営費', '5%', 'サーバー費・ドメイン等'],
    ]
)

doc.add_page_break()

# ============================================================
#  10. なぜ今・なぜ私たちか
# ============================================================
heading('10. なぜ今、なぜ私たちか')
add_table(
    ['理由', '説明'],
    [
        ['市場の追い風', '環境意識の高まりとZ世代の古着ブームで市場は拡大中'],
        ['空白領域', '古着屋に特化した地図プラットフォームは日本に存在しない'],
        ['当事者視点', '開発者自身が年数百店舗を巡る古着ファン。ユーザーの課題を肌で知っている'],
        ['低コスト展開', 'Flutter × クラウドで3プラットフォームを最小コストで運用'],
        ['スケーラビリティ', '店舗DBビジネスは規模が増えるほど価値が上がる'],
    ]
)

# ============================================================
#  フッター
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
footer_p = doc.add_paragraph('本資料に記載の数値・計画は現時点での想定であり、市場動向により変更される場合があります。')
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer_p.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_TEXT

# ===== 保存 =====
output_path = r'c:\furugi_maps\古着屋マップ_事業企画書.docx'
doc.save(output_path)
print(f'[OK] 作成完了: {output_path}')
